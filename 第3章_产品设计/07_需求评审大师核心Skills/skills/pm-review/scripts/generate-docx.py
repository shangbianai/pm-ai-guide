"""
将需求评审报告 .md 转换为 .docx
用法: python generate-docx.py <input.md> [output.docx]
字体: 宋体正文 + 黑体标题, 符合 /docx skill 格式规范
"""

import sys
import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_cell_font(cell, font_name='宋体', font_name_ascii='Times New Roman', size=Pt(9)):
    """设置表格单元格字体"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = size
            run.font.name = font_name_ascii
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_styled_paragraph(doc, text, style_name, font_name='宋体', font_name_ascii='Times New Roman',
                         size=Pt(10.5), bold=False, color=None, alignment=None,
                         space_after=Pt(6), space_before=Pt(0), line_spacing=1.3):
    """添加格式化段落"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = space_before
    pf.line_spacing = line_spacing

    run = p.add_run(text)
    run.font.name = font_name_ascii
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def add_section_heading(doc, text, level=1):
    """添加标题, 根据级别应用不同格式"""
    if level == 1:
        return add_styled_paragraph(doc, text, 'Heading 1', font_name='黑体', font_name_ascii='Arial',
                                    size=Pt(16), bold=True, space_before=Pt(24), space_after=Pt(12),
                                    line_spacing=1.15)
    elif level == 2:
        return add_styled_paragraph(doc, text, 'Heading 2', font_name='黑体', font_name_ascii='Arial',
                                    size=Pt(14), bold=True, space_before=Pt(18), space_after=Pt(8),
                                    line_spacing=1.15)
    elif level == 3:
        return add_styled_paragraph(doc, text, 'Heading 3', font_name='黑体', font_name_ascii='Arial',
                                    size=Pt(12), bold=True, space_before=Pt(12), space_after=Pt(6),
                                    line_spacing=1.15)
    else:
        return add_styled_paragraph(doc, text, 'Normal', font_name='黑体', font_name_ascii='Arial',
                                    size=Pt(12), bold=True, space_before=Pt(10), space_after=Pt(4),
                                    line_spacing=1.15)


def add_body_text(doc, text):
    """添加正文段落"""
    # 去除 markdown 格式标记
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return add_styled_paragraph(doc, text, 'Normal', line_spacing=1.3)


def add_bullet(doc, text, level=0):
    """添加项目符号列表项"""
    text = re.sub(r'^\s*[-•]\s*', '', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.3
    pf.left_indent = Cm(1.0 + level * 0.8)

    run = p.add_run('• ' + text)
    run.font.name = 'Times New Roman'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)
    return p


def add_table_from_md(doc, lines, start_idx):
    """解析 markdown 表格并添加到 docx"""
    # 找到表格的所有行 (从 start_idx 开始)
    table_lines = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        table_lines.append(lines[i].strip())
        i += 1

    if len(table_lines) < 2:
        return i

    # 解析行
    rows = []
    for line in table_lines:
        if re.match(r'^\|[\s\-:|]+\|$', line):  # 分隔行, 跳过
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)

    if not rows:
        return i

    # 创建表格
    num_cols = max(len(r) for r in rows)
    # 标准化列数
    for r in rows:
        while len(r) < num_cols:
            r.append('')

    table = doc.add_table(rows=len(rows), cols=num_cols, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            # 清除默认段落
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(cell_text)
            is_header = (row_idx == 0)
            run.bold = is_header
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            cell.paragraphs[0].paragraph_format.line_spacing = 1.15
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)

    # 添加表后空行
    doc.add_paragraph()
    return i


def add_divider(doc):
    """添加分隔线"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    # 添加底部边框作为分隔线
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def convert_md_to_docx(md_path, docx_path):
    """主转换函数"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21.0)   # A4
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.3

    # 创建封面
    for _ in range(6):
        doc.add_paragraph()

    add_styled_paragraph(doc, '需求评审报告', 'Title', font_name='黑体', font_name_ascii='Arial',
                         size=Pt(26), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_after=Pt(16), line_spacing=1.15)
    add_styled_paragraph(doc, 'AI 模拟多角色需求评审', 'Normal',
                         size=Pt(14), alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         color=RGBColor(0x66, 0x66, 0x66), space_after=Pt(40), line_spacing=1.15)
    add_styled_paragraph(doc, '产品：订单管理系统（OMS）', 'Normal',
                         size=Pt(12), alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         color=RGBColor(0x33, 0x33, 0x33), line_spacing=1.5)
    add_styled_paragraph(doc, '评审日期：2026-05-24', 'Normal',
                         size=Pt(12), alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         color=RGBColor(0x33, 0x33, 0x33), line_spacing=1.5)
    add_styled_paragraph(doc, '生成工具：pm-review Skill', 'Normal',
                         size=Pt(12), alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         color=RGBColor(0x33, 0x33, 0x33), line_spacing=1.5)
    add_styled_paragraph(doc, '文档版本：v1.0', 'Normal',
                         size=Pt(12), alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         color=RGBColor(0x33, 0x33, 0x33), line_spacing=1.5)

    doc.add_page_break()

    i = 0
    in_code_block = False
    in_table = False

    while i < len(lines):
        line = lines[i]

        # 跳过空行
        if not line.strip():
            i += 1
            continue

        # 检测代码块
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            i += 1
            continue

        # 检测表格
        if line.strip().startswith('|') and '|' in line.strip()[1:]:
            i = add_table_from_md(doc, lines, i)
            continue

        # 标题
        if line.startswith('# '):
            add_section_heading(doc, line[2:].strip(), level=1)
        elif line.startswith('## '):
            add_section_heading(doc, line[3:].strip(), level=2)
        elif line.startswith('### '):
            add_section_heading(doc, line[4:].strip(), level=3)
        elif line.startswith('#### '):
            add_section_heading(doc, line[5:].strip(), level=4)

        # 分隔线
        elif line.strip() == '---':
            add_divider(doc)

        # 引用块
        elif line.startswith('> '):
            text = line[2:].strip()
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Cm(1.0)
            pf.line_spacing = 1.3
            pf.space_after = Pt(4)
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run.italic = True

        # 项目符号
        elif re.match(r'^\s*[-•]\s+', line) or re.match(r'^\s*\d+[.\)]\s+', line):
            add_bullet(doc, line)

        # 普通文本
        else:
            # 跳过 frontmatter
            if line.strip() == '---' and i < 5:
                # 跳过到下一个 ---
                j = i + 1
                while j < len(lines) and lines[j].strip() != '---':
                    j += 1
                i = j + 1
                continue

            text = line.strip()
            if text and text != '---':
                # 处理行内格式
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                text = re.sub(r'`(.+?)`', r'\1', text)

                # 判断是否是加粗的标签行
                if text.startswith('- **') or text.startswith('**'):
                    add_bullet(doc, text)
                else:
                    add_body_text(doc, text)

        i += 1

    # 添加页脚
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run('pm-review Skill · AI 模拟评审 · 仅供内部参考')
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(docx_path)
    print(f'已生成: {docx_path}')


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('用法: python generate-docx.py <input.md> [output.docx]')
        sys.exit(1)

    md_file = sys.argv[1]
    docx_file = sys.argv[2] if len(sys.argv) > 2 else md_file.replace('.md', '.docx')

    if not os.path.exists(md_file):
        print(f'文件不存在: {md_file}')
        sys.exit(1)

    convert_md_to_docx(md_file, docx_file)
