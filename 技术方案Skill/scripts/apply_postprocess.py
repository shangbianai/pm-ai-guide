"""Post-process a pandoc-generated docx to enforce TASK-GUIDE.md formatting:

* Body paragraphs: 仿宋_GB2312 16pt, fixed 28pt line spacing, 2-char indent
  (set directly on runs — pandoc tags body as 'Body Text'/'First Paragraph',
  not 'Normal', so style-only customization would miss it)
* Image width capped to 13 cm (centered, no first-line indent)
* Tables: three-line table look (top, header bottom, bottom borders only)
  with header row 黑体 10.5pt bold, body 宋体 10.5pt
* Caption / figure title paragraphs: 宋体 9pt, centered, no indent
* Heading paragraphs: no first-line indent
* Code blocks: keep monospace but no indent
"""
import sys
from copy import deepcopy
from docx import Document
from docx.shared import Cm, Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

TARGET_IMG_WIDTH_CM = 14.5  # uniform image width (archive doc averages ~14.6cm)


def set_run_cn_font(run, ascii_font, cn_font, size_pt, bold=None, color=None):
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:cs'), ascii_font)


def clear_first_line_indent(paragraph):
    """Force first-line indent to 0 (override inheritance from Normal).

    Note: removing the w:firstLine attribute is NOT enough because the
    Normal style sets firstLine=640 and it would be re-inherited. We must
    explicitly set firstLine="0" firstLineChars="0".
    """
    pPr = paragraph._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    # Remove hanging if present (would conflict)
    for attr in ('w:hangingChars', 'w:hanging'):
        if ind.get(qn(attr)) is not None:
            del ind.attrib[qn(attr)]
    ind.set(qn('w:firstLine'), '0')
    ind.set(qn('w:firstLineChars'), '0')


def normalize_image_widths(doc):
    """Set every inline image to TARGET_IMG_WIDTH_CM wide, keep aspect ratio."""
    target_emu = Cm(TARGET_IMG_WIDTH_CM).emu
    count = 0
    drawings = doc.element.body.iter(qn('w:drawing'))
    for drawing in drawings:
        extent = drawing.find('.//' + qn('wp:extent'))
        if extent is None:
            continue
        cx = int(extent.get('cx'))
        cy = int(extent.get('cy'))
        ratio = cy / cx if cx else 0.5
        new_cx = target_emu
        new_cy = int(new_cx * ratio)
        extent.set('cx', str(new_cx))
        extent.set('cy', str(new_cy))
        for ext in drawing.findall('.//' + qn('a:ext')):
            ext.set('cx', str(new_cx))
            ext.set('cy', str(new_cy))
        count += 1
    print(f'  normalized widths on {count} images')


def center_image_paragraphs(doc):
    """Find paragraphs that contain only an image; center them, strip indent.

    Critically: explicitly write `<w:spacing w:lineRule="auto" w:line="240"/>`
    (single spacing) so images are not clipped by an inherited fixed line
    height (the body Normal style uses EXACTLY 28pt, which would crop images).
    """
    n = 0
    for p in doc.paragraphs:
        if p._element.find('.//' + qn('w:drawing')) is None:
            continue
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        clear_first_line_indent(p)
        pPr = p._element.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        # 240 = 1.0 line in twentieths-of-a-line; lineRule=auto means multiple.
        spacing.set(qn('w:line'), '240')
        spacing.set(qn('w:lineRule'), 'auto')
        spacing.set(qn('w:before'), '120')   # ~6pt
        spacing.set(qn('w:after'), '120')
        n += 1
    print(f'  centered {n} image paragraphs')


_HEADING_FONT = {
    'Heading 1': ('黑体', 16), 'Heading 2': ('楷体_GB2312', 16),
    'Heading 3': ('仿宋_GB2312', 16), 'Heading 4': ('仿宋_GB2312', 14),
    'Heading 5': ('仿宋_GB2312', 12), 'Heading 6': ('仿宋_GB2312', 12),
}


def style_headings(doc):
    """Headings: clear first-line indent AND set 黑体/楷体/仿宋 on runs directly.

    Run-level enforcement is a safety net: style-level eastAsia fonts can be
    dropped by the template's theme refs or by pandoc's style copy, which
    would leave 标题 in a default font. Setting the run font guarantees it.
    """
    n = 0
    for p in doc.paragraphs:
        sn = p.style.name
        if sn.startswith('Heading') or sn == 'Title':
            clear_first_line_indent(p)
            spec = _HEADING_FONT.get(sn)
            if spec:
                for r in p.runs:
                    set_run_cn_font(r, 'Times New Roman', spec[0], spec[1], bold=True)
            n += 1
    print(f'  cleaned indent + set font on {n} heading/title paragraphs')


# Paragraph styles that should NOT get body-text treatment.
_NON_BODY = ('Heading', 'Title', 'Caption', 'Image Caption',
             'Source Code', 'Verbatim', 'TOC', 'Table of Contents', 'Contents')
# True body paragraph styles that also get fixed line-spacing + first-line indent.
_BODY_STYLES = ('Normal', 'Body Text', 'First Paragraph', 'Compact',
                'Default Paragraph Font', 'Body')


def style_body(doc):
    """Force 正文 font/size (仿宋_GB2312 16pt) directly on every body run.

    Why directly on runs and not just via the Normal style: pandoc usually
    tags body paragraphs as 'Body Text' / 'First Paragraph', NOT 'Normal'.
    If the reference.docx only customized 'Normal', those paragraphs keep
    pandoc's default font/size — which is exactly the "字体字号不对" bug.
    Setting the run font here makes正文 correct regardless of style mapping.
    Line-spacing (固定 28pt) + 首行缩进 2 字符 are applied only to genuine
    body styles, so lists / code / captions are not wrongly indented.
    """
    n = 0
    for p in doc.paragraphs:
        sname = p.style.name or ''
        if any(sname.startswith(s) or sname == s for s in _NON_BODY):
            continue
        if p._element.find('.//' + qn('w:drawing')) is not None:
            continue  # image paragraph — handled by center_image_paragraphs
        for r in p.runs:
            set_run_cn_font(r, 'Times New Roman', '仿宋_GB2312', 16)
        if sname in _BODY_STYLES:
            pf = p.paragraph_format
            pf.line_spacing = Pt(28)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pPr = p._element.get_or_add_pPr()
            ind = pPr.find(qn('w:ind'))
            if ind is None:
                ind = OxmlElement('w:ind')
                pPr.append(ind)
            ind.set(qn('w:firstLineChars'), '200')
            ind.set(qn('w:firstLine'), '640')
        n += 1
    print(f'  styled {n} body paragraphs (仿宋 16pt)')


def _strip_table_borders(tbl):
    """Remove all borders from a table (used to un-box figure/caption frames)."""
    tblPr = tbl._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl._element.insert(0, tblPr)
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'bottom', 'left', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}'); el.set(qn('w:val'), 'nil'); borders.append(el)
    tblPr.append(borders)


def style_tables(doc):
    """Three-line table style + correct fonts for every *data* table.

    Skips figure/caption frames: pandoc (implicit_figures) wraps an image and
    its caption in a 1-cell table; a missing-image placeholder leaves an empty
    1-cell box. Styling those as 三线表 draws an ugly border around images /
    placeholders — so we strip their borders instead and don't touch fonts.
    """
    n = 0
    for tbl in doc.tables:
        has_img = tbl._element.find('.//' + qn('w:drawing')) is not None
        is_1x1 = len(tbl.rows) == 1 and len(tbl.columns) == 1
        if has_img or is_1x1:
            _strip_table_borders(tbl)
            continue
        n += 1
        tblPr = tbl._element.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl._element.insert(0, tblPr)
        # Replace borders: top + bottom only at table, plus bottom of first row
        old_borders = tblPr.find(qn('w:tblBorders'))
        if old_borders is not None:
            tblPr.remove(old_borders)
        borders = OxmlElement('w:tblBorders')
        for side in ('top', 'bottom'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '12')   # 1.5pt
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), '000000')
            borders.append(el)
        for side in ('left', 'right', 'insideV'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'nil')
            borders.append(el)
        # interior horizontal: nil — we'll add bottom on first row only
        insH = OxmlElement('w:insideH')
        insH.set(qn('w:val'), 'nil')
        borders.append(insH)
        tblPr.append(borders)
        # Header row bottom border
        rows = tbl.rows
        if rows:
            first_row = rows[0]
            for cell in first_row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                old = tcPr.find(qn('w:tcBorders'))
                if old is not None:
                    tcPr.remove(old)
                cell_borders = OxmlElement('w:tcBorders')
                bot = OxmlElement('w:bottom')
                bot.set(qn('w:val'), 'single')
                bot.set(qn('w:sz'), '8')   # 1.0pt
                bot.set(qn('w:color'), '000000')
                cell_borders.append(bot)
                tcPr.append(cell_borders)
        # Set fonts: header row = 黑体 10.5pt 加粗; body = 宋体 10.5pt
        # Looser line spacing: 22pt atLeast (≈ 1.45× of 10.5pt char height)
        for i, row in enumerate(tbl.rows):
            is_header = (i == 0)
            for cell in row.cells:
                for p in cell.paragraphs:
                    clear_first_line_indent(p)
                    pPr = p._element.get_or_add_pPr()
                    spacing = pPr.find(qn('w:spacing'))
                    if spacing is None:
                        spacing = OxmlElement('w:spacing')
                        pPr.append(spacing)
                    spacing.set(qn('w:line'), '440')         # 22pt (twentieths)
                    spacing.set(qn('w:lineRule'), 'atLeast')
                    spacing.set(qn('w:before'), '40')         # 2pt
                    spacing.set(qn('w:after'), '40')          # 2pt
                    for r in p.runs:
                        if is_header:
                            set_run_cn_font(r, 'Times New Roman', '黑体', 10.5, bold=True)
                        else:
                            set_run_cn_font(r, 'Times New Roman', '宋体', 10.5)
                    if not p.runs and p.text == '':
                        # ensure cell has at least empty styled run
                        r = p.add_run('')
                        if is_header:
                            set_run_cn_font(r, 'Times New Roman', '黑体', 10.5, bold=True)
                        else:
                            set_run_cn_font(r, 'Times New Roman', '宋体', 10.5)
        # Table layout: auto-fit content
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    print(f'  styled {n} tables (three-line)')


def style_captions(doc):
    """Pandoc's captions land in style 'Image Caption' or 'Caption'."""
    n = 0
    for p in doc.paragraphs:
        if p.style.name in ('Image Caption', 'Caption'):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            clear_first_line_indent(p)
            pf = p.paragraph_format
            pf.line_spacing = Pt(14)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.space_before = Pt(2)
            pf.space_after = Pt(6)
            for r in p.runs:
                set_run_cn_font(r, 'Times New Roman', '宋体', 9)
            n += 1
    print(f'  styled {n} captions')


def main(path_in, path_out):
    doc = Document(path_in)
    print(f'Loaded {path_in}')
    normalize_image_widths(doc)
    center_image_paragraphs(doc)
    style_headings(doc)
    style_body(doc)
    style_tables(doc)
    style_captions(doc)
    doc.save(path_out)
    print(f'Saved {path_out}')


if __name__ == '__main__':
    main(sys.argv[1], sys.argv[2])
